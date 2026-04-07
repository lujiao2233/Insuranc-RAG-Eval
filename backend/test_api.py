"""API接口测试脚本
测试认证系统和文档分析模块的API接口
"""
import httpx
import asyncio
import os
import sys
import json
import time
from pathlib import Path
from io import BytesIO

BASE_URL = "http://localhost:8000"


def create_test_pdf():
    """创建测试PDF文件"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.setFont("Helvetica", 12)
        c.drawString(100, 750, "Test Document for API Testing")
        c.drawString(100, 720, "This is a test PDF document.")
        c.drawString(100, 690, "Content includes:")
        c.drawString(120, 660, "- Document analysis testing")
        c.drawString(120, 630, "- API integration verification")
        c.drawString(120, 600, "- Authentication flow testing")
        c.save()
        buffer.seek(0)
        return buffer
    except ImportError:
        return None


def create_test_docx():
    """创建测试DOCX文件"""
    try:
        from docx import Document
        
        buffer = BytesIO()
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test DOCX document for API testing.')
        doc.add_paragraph('Content includes:')
        doc.add_paragraph('Document analysis testing', style='List Bullet')
        doc.add_paragraph('API integration verification', style='List Bullet')
        doc.add_paragraph('Authentication flow testing', style='List Bullet')
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except ImportError:
        return None


class APITester:
    def __init__(self):
        self.client = httpx.AsyncClient(timeout=30.0)
        self.token = None
        self.user_id = None
        self.document_id = None
        self.username = f"testuser_{int(time.time())}"
    
    async def close(self):
        await self.client.aclose()
    
    def get_headers(self, with_auth=False):
        headers = {"Content-Type": "application/json"}
        if with_auth and self.token:
            headers["Authorization"] = f"Bearer {self.token}"
        return headers
    
    async def test_health(self):
        """测试健康检查接口"""
        print("\n" + "=" * 50)
        print("测试健康检查接口...")
        
        try:
            response = await self.client.get(f"{BASE_URL}/health")
            print(f"  GET /health")
            print(f"  状态码: {response.status_code}")
            print(f"  响应: {response.json()}")
            return response.status_code == 200
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_auth_register(self):
        """测试用户注册接口"""
        print("\n" + "=" * 50)
        print("测试认证API - 用户注册...")
        
        try:
            response = await self.client.post(
                f"{BASE_URL}/api/v1/auth/register",
                json={
                    "username": self.username,
                    "email": f"{self.username}@test.com",
                    "password": "testpassword123",
                    "full_name": "Test User"
                }
            )
            print(f"  POST /api/v1/auth/register")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 201:
                data = response.json()
                self.user_id = data.get("id")
                print(f"  用户ID: {self.user_id}")
                print(f"  用户名: {data.get('username')}")
                print("  ✓ 注册成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_auth_login(self):
        """测试用户登录接口"""
        print("\n" + "=" * 50)
        print("测试认证API - 用户登录...")
        
        try:
            response = await self.client.post(
                f"{BASE_URL}/api/v1/auth/login",
                data={
                    "username": self.username,
                    "password": "testpassword123"
                },
                headers={"Content-Type": "application/x-www-form-urlencoded"}
            )
            print(f"  POST /api/v1/auth/login")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                self.token = data.get("access_token")
                print(f"  Token: {self.token[:50]}..." if self.token else "  Token: None")
                print("  ✓ 登录成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_auth_me(self):
        """测试获取当前用户信息接口"""
        print("\n" + "=" * 50)
        print("测试认证API - 获取当前用户信息...")
        
        if not self.token:
            print("  跳过: 未登录")
            return False
        
        try:
            response = await self.client.get(
                f"{BASE_URL}/api/v1/auth/me",
                headers=self.get_headers(with_auth=True)
            )
            print(f"  GET /api/v1/auth/me")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                print(f"  用户名: {data.get('username')}")
                print(f"  邮箱: {data.get('email')}")
                print("  ✓ 获取用户信息成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_auth_refresh_token(self):
        """测试刷新Token接口"""
        print("\n" + "=" * 50)
        print("测试认证API - 刷新Token...")
        
        if not self.token:
            print("  跳过: 未登录")
            return False
        
        try:
            response = await self.client.post(
                f"{BASE_URL}/api/v1/auth/refresh-token",
                headers=self.get_headers(with_auth=True)
            )
            print(f"  POST /api/v1/auth/refresh-token")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                new_token = data.get("access_token")
                print(f"  新Token: {new_token[:50]}...")
                print("  ✓ 刷新Token成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_document_list(self):
        """测试文档列表接口"""
        print("\n" + "=" * 50)
        print("测试文档API - 获取文档列表...")
        
        try:
            response = await self.client.get(
                f"{BASE_URL}/api/v1/documents/",
                headers=self.get_headers(with_auth=True)
            )
            print(f"  GET /api/v1/documents/")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                print(f"  文档总数: {data.get('total', 0)}")
                print(f"  当前页: {data.get('page', 1)}")
                print("  ✓ 获取列表成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_document_upload_pdf(self):
        """测试PDF文档上传接口"""
        print("\n" + "=" * 50)
        print("测试文档API - PDF文档上传...")
        
        pdf_buffer = create_test_pdf()
        if not pdf_buffer:
            print("  跳过: 无法创建测试PDF（缺少reportlab库）")
            return False
        
        try:
            files = {
                "file": ("test_document.pdf", pdf_buffer, "application/pdf")
            }
            headers = {}
            if self.token:
                headers["Authorization"] = f"Bearer {self.token}"
            
            response = await self.client.post(
                f"{BASE_URL}/api/v1/documents/upload",
                files=files,
                headers=headers
            )
            print(f"  POST /api/v1/documents/upload")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                self.document_id = data.get("id")
                print(f"  文档ID: {self.document_id}")
                print(f"  文件名: {data.get('filename')}")
                print(f"  文件类型: {data.get('file_type')}")
                print("  ✓ PDF上传成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_document_upload_docx(self):
        """测试DOCX文档上传接口"""
        print("\n" + "=" * 50)
        print("测试文档API - DOCX文档上传...")
        
        docx_buffer = create_test_docx()
        if not docx_buffer:
            print("  跳过: 无法创建测试DOCX")
            return False
        
        try:
            files = {
                "file": ("test_document.docx", docx_buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            }
            headers = {}
            if self.token:
                headers["Authorization"] = f"Bearer {self.token}"
            
            response = await self.client.post(
                f"{BASE_URL}/api/v1/documents/upload",
                files=files,
                headers=headers
            )
            print(f"  POST /api/v1/documents/upload")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                if not self.document_id:
                    self.document_id = data.get("id")
                print(f"  文档ID: {data.get('id')}")
                print(f"  文件名: {data.get('filename')}")
                print(f"  文件类型: {data.get('file_type')}")
                print("  ✓ DOCX上传成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_document_get(self):
        """测试获取单个文档接口"""
        print("\n" + "=" * 50)
        print("测试文档API - 获取单个文档...")
        
        if not self.document_id:
            print("  跳过: 没有文档ID")
            return False
        
        try:
            response = await self.client.get(
                f"{BASE_URL}/api/v1/documents/{self.document_id}",
                headers=self.get_headers(with_auth=True)
            )
            print(f"  GET /api/v1/documents/{self.document_id}")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                print(f"  文件名: {data.get('filename')}")
                print(f"  文件类型: {data.get('file_type')}")
                print(f"  是否已分析: {data.get('is_analyzed')}")
                print("  ✓ 获取文档成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_document_stats(self):
        """测试文档统计接口"""
        print("\n" + "=" * 50)
        print("测试文档API - 获取文档统计...")
        
        try:
            response = await self.client.get(
                f"{BASE_URL}/api/v1/documents/stats/summary",
                headers=self.get_headers(with_auth=True)
            )
            print(f"  GET /api/v1/documents/stats/summary")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                print(f"  总文档数: {data.get('total_documents', 0)}")
                print(f"  已分析: {data.get('analyzed_documents', 0)}")
                print(f"  未分析: {data.get('unanalyzed_documents', 0)}")
                if data.get('type_distribution'):
                    print(f"  类型分布: {data.get('type_distribution')}")
                print("  ✓ 获取统计成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_analysis_status(self):
        """测试分析状态接口"""
        print("\n" + "=" * 50)
        print("测试分析API - 获取分析状态...")
        
        if not self.document_id:
            print("  跳过: 没有文档ID")
            return False
        
        try:
            response = await self.client.get(
                f"{BASE_URL}/api/v1/analysis/status/{self.document_id}",
                headers=self.get_headers(with_auth=True)
            )
            print(f"  GET /api/v1/analysis/status/{self.document_id}")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                print(f"  是否已分析: {data.get('is_analyzed')}")
                print(f"  状态: {data.get('status')}")
                print(f"  有元数据: {data.get('has_metadata')}")
                print(f"  有大纲: {data.get('has_outline')}")
                print("  ✓ 获取状态成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_analysis_trigger(self):
        """测试触发文档分析接口"""
        print("\n" + "=" * 50)
        print("测试分析API - 触发文档分析...")
        
        if not self.document_id:
            print("  跳过: 没有文档ID")
            return False
        
        try:
            response = await self.client.post(
                f"{BASE_URL}/api/v1/analysis/analyze/{self.document_id}",
                headers=self.get_headers(with_auth=True)
            )
            print(f"  POST /api/v1/analysis/analyze/{self.document_id}")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                print(f"  消息: {data.get('message')}")
                print(f"  状态: {data.get('status')}")
                print("  ✓ 触发分析成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def test_auth_logout(self):
        """测试用户登出接口"""
        print("\n" + "=" * 50)
        print("测试认证API - 用户登出...")
        
        if not self.token:
            print("  跳过: 未登录")
            return False
        
        try:
            response = await self.client.post(
                f"{BASE_URL}/api/v1/auth/logout",
                headers=self.get_headers(with_auth=True)
            )
            print(f"  POST /api/v1/auth/logout")
            print(f"  状态码: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                print(f"  消息: {data.get('message')}")
                print("  ✓ 登出成功")
                return True
            else:
                print(f"  响应: {response.text}")
                return False
        except Exception as e:
            print(f"  错误: {e}")
            return False
    
    async def run_all_tests(self):
        """运行所有测试"""
        print("\n" + "=" * 60)
        print("RAG评估系统 - API接口测试")
        print("=" * 60)
        
        results = {}
        
        # 认证测试
        results["健康检查"] = await self.test_health()
        results["用户注册"] = await self.test_auth_register()
        results["用户登录"] = await self.test_auth_login()
        results["获取用户信息"] = await self.test_auth_me()
        results["刷新Token"] = await self.test_auth_refresh_token()
        
        # 文档测试
        results["文档列表"] = await self.test_document_list()
        results["PDF上传"] = await self.test_document_upload_pdf()
        results["DOCX上传"] = await self.test_document_upload_docx()
        results["获取文档"] = await self.test_document_get()
        results["文档统计"] = await self.test_document_stats()
        
        # 分析测试
        results["分析状态"] = await self.test_analysis_status()
        results["触发分析"] = await self.test_analysis_trigger()
        
        # 再次检查分析状态
        await asyncio.sleep(1)
        results["分析状态(触发后)"] = await self.test_analysis_status()
        
        # 登出测试
        results["用户登出"] = await self.test_auth_logout()
        
        print("\n" + "=" * 60)
        print("测试结果汇总:")
        print("=" * 60)
        
        all_passed = True
        for name, passed in results.items():
            status = "✓ 通过" if passed else "✗ 失败"
            print(f"  {name}: {status}")
            if not passed:
                all_passed = False
        
        print("\n" + "=" * 60)
        if all_passed:
            print("所有测试通过!")
        else:
            print("部分测试失败，请检查上述错误信息。")
        print("=" * 60)
        
        return all_passed


async def main():
    tester = APITester()
    try:
        success = await tester.run_all_tests()
        sys.exit(0 if success else 1)
    finally:
        await tester.close()


if __name__ == "__main__":
    asyncio.run(main())
